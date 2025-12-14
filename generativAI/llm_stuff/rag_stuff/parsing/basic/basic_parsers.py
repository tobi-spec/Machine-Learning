import os
from typing import List
from langchain_core.documents import Document
from pathlib import Path
import pandas as pd
import pypdf
import docx

def resolve_relative_path(path: str) -> Path:
    base_dir = Path(__file__).resolve().parent.parent.parent / "sample_documents"
    return base_dir / path


class BasicPDFParser:
    def read(self, pdf_path: str) -> List[Document]:
        documents = []
        reader = pypdf.PdfReader(resolve_relative_path(pdf_path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            documents.append(Document(page_content=text, metadata={"source": pdf_path, "page": i+1, "type": "pdf"}))
        return documents


class BasicWordParser:
    def read(self, word_path: str) -> List[Document]:
        documents = []
        doc = docx.Document(str(resolve_relative_path(word_path)))
        text = "\n".join(p.text for p in doc.paragraphs)
        documents.append(Document(
            page_content=text,
            metadata={"source": word_path, "type": "word"}
        ))
        return documents


class ExcelParser:
    def read(self, excel_path: str) -> List[Document]:
        documents = []
        xls = pd.ExcelFile(resolve_relative_path(excel_path))
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            for idx, row in df.iterrows():
                row_text = ",".join(f"{col}: {row[col]}" for col in df.columns)
                documents.append(Document(page_content=row_text, metadata={
                        "source": excel_path,
                        "type": "excel",
                        "sheet": sheet_name,
                        "row_index": int(idx)
                    }))
        return documents

